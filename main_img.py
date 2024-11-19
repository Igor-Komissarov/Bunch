import pandas as pd
import yaml
import re
from collections import Counter
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import shutil
from googletrans import Translator



# Функция для сохранения графиков как изображений
def save_chart_as_image(fig, output_path, width=1920, height=1080):
    fig.update_layout(width=width, height=height)
    fig.write_image(output_path)

# Функция для построения графика "Динамика патентования компании"
def grapth(yaml_file, bunch_file, output_image):
    with open(yaml_file, 'r', encoding='utf-8') as file:
        columns = yaml.safe_load(file)

    df = pd.read_excel(bunch_file, sheet_name='SHEET')

    def extract_year_counts(df, column_name):
        valid_entries = df[column_name].dropna()
        years = valid_entries[valid_entries != 'no data'].str[:4]
        year_counts = Counter(years)
        return pd.DataFrame(list(year_counts.items()), columns=['Year', column_name.split()[1].capitalize()])

    priority_df = extract_year_counts(df, columns['priority_date_column'])
    publication_df = extract_year_counts(df, columns['publication_date_column'])
    grant_df = extract_year_counts(df, columns['grant_date_column'])

    merged_df = priority_df.merge(publication_df, on='Year', how='outer').merge(grant_df, on='Year', how='outer').fillna(0)

    merged_df['Year'] = merged_df['Year'].astype(int)
    merged_df = merged_df[(merged_df['Year'] >= 2014) & (merged_df['Year'] <= 2023)].sort_values(by='Year')

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=merged_df['Year'], y=merged_df['Priority'], mode='lines+markers', name='Патентные семейства',
        line=dict(color='rgba(102, 153, 204, 0.8)', width=6, shape='spline'),
        hovertemplate='<b>Год:</b> %{x}<br><b>Количество:</b> %{y}<extra></extra>'
    ))
    fig.add_trace(go.Scatter(
        x=merged_df['Year'], y=merged_df['Publication'], mode='lines+markers', name='Патентные публикации',
        line=dict(color='rgba(153, 204, 102, 0.8)', width=6, shape='spline'),
        hovertemplate='<b>Год:</b> %{x}<br><b>Количество:</b> %{y}<extra></extra>'
    ))
    fig.add_trace(go.Scatter(
        x=merged_df['Year'], y=merged_df['Grant'], mode='lines+markers', name='Патенты',
        line=dict(color='rgba(255, 153, 102, 0.8)', width=6, shape='spline'),
        hovertemplate='<b>Год:</b> %{x}<br><b>Количество:</b> %{y}<extra></extra>'
    ))

    y_max = max(merged_df[['Priority', 'Publication', 'Grant']].max()) + 1
    #y_ticks = list(range(0, int(y_max) + 1))  # Генерируем значения для оси Y от 0 до максимума

    fig.update_layout(
        title='Динамика патентования компании',
        xaxis=dict(
            title='Год',
            tickmode='linear',
            dtick=1,
            range=[2013.5, 2023.5],  # Отдаление оси X
            title_standoff=25,
            tickfont=dict(size=29),
            automargin=True
        ),
        yaxis=dict(
            title='Количество',
            range=[-1, y_max],
            # tickvals=y_ticks,  # Указываем только положительные значения
            # ticktext=[str(val) for val in y_ticks],  # Преобразуем значения в текст
            title_standoff=25,
            tickfont=dict(size=29),
            automargin=True
        ),
        template='plotly_white',
        font=dict(family='Arial, sans-serif', size=28, color='black'),
        legend=dict(
            title=dict(text='Категории<br>', font=dict(size=28)),
            font=dict(family='Arial, sans-serif', size=28, color='black'),
            orientation='v',
            y=0.6,
            x=1.05,
            xanchor='left',
            itemclick='toggleothers',
            itemsizing='constant',
            traceorder='normal',
            itemwidth=50
        ),
        margin=dict(t=80, b=40, l=80, r=20)
    )

    save_chart_as_image(fig, output_image)
    print(f"График сохранен как {output_image}")

# Функция для создания круговых диаграмм
# Функция для создания круговых диаграмм (шайб)
def create_legal_status_pie_charts(file_path, output_image_1, output_image_2):
    """
    Создаёт два круговых графика (шайбы): 
    1) для столбца 'Legal state (Alive, Dead)' 
    2) для столбца 'Legal status (Pending, Granted, Revoked, Expired, Lapsed)' 
    с правильными цветами, шрифтами, переводом и легендой.
    """
    df = pd.read_excel(file_path)

    # Внутренняя функция для обработки статусов
    def process_legal_status(df, column_name, values):
        status_counts = df[column_name].dropna().apply(lambda x: re.findall(r'\b(?:' + '|'.join(values) + r')\b', x))
        status_counts = status_counts.explode().value_counts().reset_index()
        status_counts.columns = ['Status', 'Count']
        return status_counts

    # Внутренняя функция для построения и сохранения шайб
    def plot_pie_chart(data, title, output_path, status_labels, colors, width=1920, height=1080):
        data['Status'] = data['Status'].replace(status_labels)
        fig = px.pie(data, names='Status', values='Count', title=title, color='Status', color_discrete_map=colors)

        fig.update_traces(
            textinfo='percent+label',
            textposition='outside'
        )
        fig.update_layout(
            font=dict(family='Arial, sans-serif', size=26, color='black'),
            legend=dict(
                title=dict(text='Категории', font=dict(size=20)),
                font=dict(size=24),
                orientation='v',
                x=100000,  # Крайний правый угол
                y=11110,
                xanchor='left',
                yanchor='middle'
            ),
            margin=dict(t=80, b=40, l=80, r=40)
        )
        fig.update_layout(width=width, height=height)
        fig.write_image(output_path)
        print(f"График '{title}' сохранён как {output_path}")

    # Построение графика для Legal state (Alive, Dead)
    legal_state_values = ['ALIVE', 'DEAD']
    legal_state_data = process_legal_status(df, 'Legal state (Alive, Dead)', legal_state_values)
    legal_state_labels = {'ALIVE': 'Действующие патентные семейства', 'DEAD': 'Недействующие патентные семейства'}
    legal_state_colors = {'Действующие патентные семейства': '#4682B4', 'Недействующие патентные семейства': 'red'}
    plot_pie_chart(
        legal_state_data, 
        'Правовое состояние с учётом делопроизводства', 
        output_image_1, 
        legal_state_labels, 
        legal_state_colors
    )

    # Построение графика для Legal status (Pending, Granted, Revoked, Expired, Lapsed)
    legal_status_values = ['PENDING', 'GRANTED', 'REVOKED', 'EXPIRED', 'LAPSED']
    legal_status_data = process_legal_status(df, 'Legal status (Pending, Granted, Revoked, Expired, Lapsed)', legal_status_values)
    legal_status_labels = {
        'PENDING': 'Заявки на рассмотрении',
        'GRANTED': 'Действующие патенты',
        'REVOKED': 'Отозванные заявки',
        'EXPIRED': 'Патенты с истёкшим сроком действия',
        'LAPSED': 'Патенты, прекратившие действие по иным причинам'
    }
    legal_status_colors = {
        'Заявки на рассмотрении': '#B3D7D4',
        'Действующие патенты': '#4682B4',
        'Отозванные заявки': '#8A2BE2',
        'Патенты с истёкшим сроком действия': '#32CD32',
        'Патенты, прекратившие действие по иным причинам': 'red'
    }
    plot_pie_chart(
        legal_status_data, 
        'Правовой статус с учётом делопроизводства', 
        output_image_2, 
        legal_status_labels, 
        legal_status_colors
    )

# Полный код с остальными функциями.
# Функция для создания рейтинга стран
def country_rank(patent_data_path, country_mapping_path, output_path):
    df_patents = pd.read_excel(patent_data_path)
    df_country_mapping = pd.read_excel(country_mapping_path)

    country_dict = dict(zip(df_country_mapping['Country Code'], df_country_mapping['True Russian Country Name']))
    countries = df_patents['Страны (юрисдикции) патентования'].dropna().str.split(', ').explode()
    filtered_countries = countries.map(country_dict).dropna()

    country_counts = filtered_countries.value_counts().head(20).reset_index()
    country_counts.columns = ['Страна (юрисдикция)', 'Число патентных семейств']

    fig = px.bar(
        country_counts,
        x='Страна (юрисдикция)',
        y='Число патентных семейств',
        title='Рейтинг стран (юрисдикций)',
        text='Число патентных семейств'
    )
    fig.update_traces(
        marker_color='royalblue',
        textposition='outside'
    )
    fig.update_layout(
        xaxis_title='Страна (юрисдикция)',
        yaxis_title='Число патентных семейств',
        font=dict(family='Arial, sans-serif', size=26, color='black'),
        xaxis=dict(title_standoff=25, tickfont=dict(size=29), automargin=True),
        yaxis=dict(title_standoff=25, tickfont=dict(size=29), automargin=True),
        margin=dict(t=80, b=40, l=80, r=40)
    )
    save_chart_as_image(fig, output_path)
    print(f"График стран сохранен как {output_path}")

# Функция для создания рейтинга изобретателей
def inventors_rank(patent_data_path, output_path):
    df_patents = pd.read_excel(patent_data_path)
    inventor_column = 'Inventors'
    df_patents[inventor_column] = df_patents[inventor_column].dropna().str.split('\n')
    df_exploded = df_patents.explode(inventor_column)

    inventor_counts = df_exploded.groupby(inventor_column).size().nlargest(10).reset_index(name='Число патентных семейств')
    fig = px.bar(
        inventor_counts,
        x='Inventors',
        y='Число патентных семейств',
        title='Рейтинг авторов изобретений',
        text='Число патентных семейств'
    )
    fig.update_traces(
        marker_color='darkorange',
        textposition='outside'
    )
    fig.update_layout(
        xaxis_title='Автор изобретения',
        yaxis_title='Число патентных семейств',
        font=dict(family='Arial, sans-serif', size=26, color='black'),
        xaxis=dict(title_standoff=25, tickfont=dict(size=29), automargin=True),
        yaxis=dict(title_standoff=25, tickfont=dict(size=29), automargin=True),
        margin=dict(t=80, b=40, l=80, r=40)
    )
    save_chart_as_image(fig, output_path)
    print(f"График авторов сохранен как {output_path}")

# Функция для создания таблицы с топ-5 патентами
def create_top_patent_word_report_with_translation(file_path_1, file_path_2, output_word_file):
    """
    Загружает данные из двух файлов Excel, объединяет, переводит Title на русский для топ-5 записей,
    сортирует по 'Patent Strength' и сохраняет их в новый Word файл в нужной последовательности.
    """
    
    def clean_text(value):
        """Очищает текст, заменяя начальные символы новой строки (\n) на точку с запятой (;)."""
        if isinstance(value, str):
            return value.replace('\n', ';\n').strip()
        return value

    # Инициализация переводчика
    translator = Translator()

    # Загрузка данных из первого файла
    df1 = pd.read_excel(file_path_1)
    df2 = pd.read_excel(file_path_2)

    # Переименовать нужные столбцы в первом файле
    columns_mapping = {'Название на английском языке': 'Title', 'Сила патентного семейства': 'Patent strength'}
    df1_selected = df1[list(columns_mapping.keys())].rename(columns=columns_mapping)

    # Выбираем нужные столбцы из второго файла
    additional_columns = {'Current assignees': 'Current Assignees', 'Earliest publication number': 'Earliest Publication Number'}
    df2_selected = df2[list(additional_columns.keys())].rename(columns=additional_columns)

    # Объединяем данные из двух файлов
    combined_df = pd.concat([df2_selected, df1_selected], axis=1)

    # Очистка строк от начальных \n
    for col in combined_df.columns:
        combined_df[col] = combined_df[col].apply(clean_text)

    # Сортировка по "Patent strength" и выбор топ-5 записей
    top_5_df = combined_df.sort_values(by='Patent strength', ascending=False).head(5)

    # Перевод столбца Title на русский для топ-5
    top_5_df['Title (Russian)'] = top_5_df['Title'].apply(lambda x: translator.translate(x, src='en', dest='ru').text)

    # Переименование и установка нужного порядка столбцов
    top_5_df.insert(0, '№', range(1, 6))
    top_5_df = top_5_df.rename(columns={
        'Earliest Publication Number': 'Номер публикации',
        'Current Assignees': 'Название компании',
        'Title': 'Название изобретения на английском',
        'Title (Russian)': 'Название изобретения на русском',
        'Patent strength': 'Сила патента'
    })
    top_5_df = top_5_df[['№', 'Номер публикации', 'Название компании', 'Название изобретения на английском', 'Название изобретения на русском', 'Сила патента']]

    # Создание документа Word
    doc = Document()
    doc.add_heading('Топ-5 патентов по силе патента', level=1)

    # Добавление таблицы
    table = doc.add_table(rows=1, cols=len(top_5_df.columns))
    table.style = 'Table Grid'

    # Заполнение заголовков таблицы
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(top_5_df.columns):
        hdr_cells[i].text = column_name

    # Заполнение данных таблицы
    for index, row in top_5_df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Форматирование шрифта для таблицы
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True  # Заголовок таблицы полужирный
                run.font.size = Pt(12)  # Размер шрифта заголовка

    # Сохранение документа Word
    doc.save(output_word_file)
    print(f"Топ-5 записей по Patent Strength сохранены в Word файл '{output_word_file}'")


# Функция для создания отчета



# Основной код для выполнения
company = 'Beihang'
yaml_file = '/Users/igorkomissarov/Bunch/column_names.yaml'
bunch_file = f'/Users/igorkomissarov/Bunch/Company/{company}/Diversity bunch {company}.xlsx'
output_dir = f'/Users/igorkomissarov/Bunch/Company/{company}'
img_dir = f'{output_dir}/Img'
country_mapping_path = '/Users/igorkomissarov/Bunch/Разное/Расшифровка двухбуквенных кодов юрисдикций.xlsx'
file_path_1 = f'/Users/igorkomissarov/Bunch/Company/{company}/Restore_{company}.xlsx'

graph_output_1 = f'{img_dir}/graph_1.png'
graph_output_2 = f'{img_dir}/country_rank.png'
graph_output_3 = f'{img_dir}/legal_state_distribution.png'
graph_output_4 = f'{img_dir}/legal_status_distribution.png'
graph_output_5 = f'{img_dir}/top_inventor_rank.png'

table_output_word = f'{output_dir}/top_5_combined_patents_by_strength_with_translation.docx'

if os.path.exists(img_dir):
    shutil.rmtree(img_dir)
os.makedirs(img_dir)

# Генерация данных
grapth(yaml_file, bunch_file, graph_output_1)
create_top_patent_word_report_with_translation(file_path_1, bunch_file, table_output_word)
country_rank(file_path_1, country_mapping_path, graph_output_2)
inventors_rank(bunch_file, graph_output_5)
create_legal_status_pie_charts(bunch_file, graph_output_3, graph_output_4)

print("Все этапы завершены успешно!")
